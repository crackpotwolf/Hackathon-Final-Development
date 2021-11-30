from docx import Document, table


class DocxLoader:
  """Класс обработки файлов формата *.docx
  """
  
  def __init__(self):
    """Метод инициализации объекта класса
    """
    self.data = {}
  
  def __read_file__(self, filepath: str):
    """Метод считывания данных из файла

    Args:
        filepath (str): [путь к файлу]
    """
    def get_data_from_table(table: table):
      """Метод считывания данных из таблицы

      Args:
          table (table): [текущая таблицы]

      Returns:
          [List]: [лист словарей, содержащие данные по ячейкам в формате название столбца: значение ячейки]
      """
      # метод парсинга таблицы
      data = []
      keys = None
      # проход по всем строкам таблицы
      for i, row in enumerate(table.rows):
        # забрать текст из ячейки
        text = (cell.text for cell in row.cells)
        # если строка первая (0), то  задать заголовки столбцов
        if i == 0:
            keys = tuple(text)
            continue
        # запаковывание данных ячейки
        row_data = dict(zip(keys, text))
        data.append(row_data)
      return data
    
    def get_data_from_tables(document: Document):
      """Метод определения таблиц в документе для их дальнейшего парсинга

      Args:
          document (Document): [открытый документ]

      Returns:
          [Dict]: [словарь данных таблиц в формате номер таблицы по счету: данные из таблицы]
      """
      # получить все таблицы из документа
      tables = document.tables
      # собрать данные из таблиц по одной
      tables_data = {}
      for num_table, current_table in enumerate(tables):
        # получить данные из текущей таблицы
        current_table_data = get_data_from_table(table=current_table)
        # сохранить спарсенные из таблицы данные под ключом с номером таблицы
        tables_data["table_%s" % str(num_table + 1)] = current_table_data
      return tables_data
      
    def get_full_text(document: Document):
      """Метод считываня текста из документа

      Args:
          document (Document): [открытый документ]

      Returns:
          [Dict]: [словарь с текстом из документа]
      """
      # получить все абзацы из документа
      paragraphs = document.paragraphs
      full_text = ""
      # получить текст из каждого найденного параграфа
      for current_paragraph in paragraphs:
        full_text += current_paragraph.text
      return full_text
    
    # открыть документ
    document = Document(docx=filepath)
    # получить данные из таблиц
    data_from_tables = get_data_from_tables(document=document)
    if len(data_from_tables) > 0:
      self.data = data_from_tables
    # получить текст
    self.data["full_text"] = get_full_text(document=document)
    
